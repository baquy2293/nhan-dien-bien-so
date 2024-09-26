import os
import cv2
from ultralytics import YOLO
from docx import Document
def process_image(model, image, source):
    mytext = ""
    results = model(image)
    boxes = results[0].boxes
    image_model = YOLO('best1.pt')

    if boxes.shape[0] > 0:
        x_min = int(boxes.xyxy[0][0])
        y_min = int(boxes.xyxy[0][1])
        x_max = int(boxes.xyxy[0][2])
        y_max = int(boxes.xyxy[0][3])
        cropped_image = image[y_min:y_max, x_min:x_max]
        results = image_model.predict(cropped_image)
        if source == 'pi':
            cv2.imwrite("static/cropped_image_out.jpg", cropped_image)
        elif source == 'ras':
            cv2.imwrite("static/cropped_image_in.jpg", cropped_image)
            
        gray_image = cv2.cvtColor(cropped_image, cv2.COLOR_BGR2GRAY)
        image_results = image_model(cropped_image)
        # print(image_results)
        # Extract bounding boxes, classes, names, and confidences
        
        boxes = results[0].boxes.xyxy.tolist()
        classes = results[0].boxes.cls.tolist()
        names = results[0].names
        confidences = results[0].boxes.conf.tolist()
        a = []
        for _ in range(15):
            hang = [0] * 3  # Create a row containing 3 elements with value 0
            a.append(hang)

        row = 0 
        for box, cls, conf in zip(boxes, classes, confidences):
            x1, y1, x2, y2 = box
            confidence = conf
            detected_class = cls
            name = names[int(cls)]
            a[row][0] = x1
            a[row][1] = y1
            a[row][2] = detected_class
            row += 1
        # for i in range(row):
        #     print(f"{a[i][0]} x {a[i][1]} x {a[i][2]}")
        # print(f"row: {row}")
        line1 = [15]*row
        line2 = [15]*row
        visit = [15]*row
        for i in range(row): 
            line1[i] = 10000
            line2[i] = 10000
            visit[i] = 0
                
        check = 0
        for i in range(row): 
            if(a[i][1]<30):
                if(visit[i] == 0):
                    check = i
                    for j in range(row):
                        if(visit[j] == 0):
                            visit[j] = 1
                            # print(f"{a[check][1]} x {a[j][1]}")
                            if(a[j][1] - a[check][1] <= 20):
                                line1.append(j)
                            else:
                                line2.append(j)
            
            
        # print(line1)
        # print(line2)
        for i in range(len(line1)-1):
            for j in range(i + 1, len(line1)):
                if(line1[i] != 10000):
                    if a[line1[i]][0] > a[line1[j]][0]:
                        tem = line1[i]
                        line1[i] = line1[j]
                        line1[j] = tem
                        
        for i in range(len(line2)-1):
            for j in range(i + 1, len(line2)):
                if(line2[i] != 10000):
                    if a[line2[i]][0] > a[line2[j]][0]:
                        # Swap elements if condition is true
                        tem = line2[i]
                        line2[i] = line2[j]
                        line2[j] = tem
                        
        
        for i in range(len(line1)):
            if(line1[i] != 10000):
                mytext += names[a[line1[i]][2]]
        for i in range(len(line2)):
            if(line2[i] != 10000):
                mytext += names[a[line2[i]][2]]
    #     print(mytext)
    #     print(line1)
    #     print(line2)
    # print(mytext)
    return mytext

# Path to the directory containing images
image_dir = "C://Users//anh//Downloads//Bien_sach"

# Initialize the model
model = YOLO('best.pt')

# Source type (pi or ras)
source = 'pi'  # or 'ras'

# List all image files in the directory
image_paths = [os.path.join(image_dir, f) for f in os.listdir(image_dir) if f.endswith(('jpg', 'jpeg', 'png'))]

# Create a Word document
doc = Document()
doc.add_heading('License Plate Recognition Results', 0)

# Process each image and save the results in the Word document
for image_path in image_paths:
    image = cv2.imread(image_path)
    result = process_image(model, image, source)
    doc.add_paragraph(f"Processed {image_path}: {result}")

# Save the Word document
doc.save('License_Plate_Recognition_Results.docx')
